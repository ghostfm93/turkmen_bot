import os
from sys import exc_info
from docxtpl import DocxTemplate
import json
from openpyxl import load_workbook
from docx import Document
from docx2pdf import convert
from datetime import datetime, timedelta
import pythoncom
user_data = {}
class MakeDocs:
    def __init__(self, context:dict):
        self.context = context
        self.contract_filename = ''
        self.petition_filename = ''
        self.filepaths = []
        self.months = {
                        '01': 'января',
                        '02': 'февраля',
                        '03': 'марта',
                        '04': 'апреля',
                        '05': 'мая',
                        '06': 'июня',
                        '07': 'июля',
                        '08': 'августа',
                        '09': 'сентября',
                        '10': 'октября',
                        '11': 'ноября',
                        '12': 'декабря'
        }

    @staticmethod
    def read_contracts_data():
        '''читаем номера договоров из файла'''
        with open('settings/contracts.json', 'r') as js_file:
            contracts_data = json.load(js_file)
            print('Считали данные о договорах из файла')
            return contracts_data

    def write_contracts_data(self, data):
        '''записываем изменения номеров договоров в файл'''
        with open('settings/contracts.json', 'w') as js_file:
            json.dump(data, js_file)
            print('Записали данные о договорах в файл')

    def make_petition(self, data):
        '''делаем ходатайство'''
        # добавляем в контекст номер ходатайства и делаем манипуляции с ФИО
        self.context.update({'petition_number': data['petition_number']})
        self.context.update({'first_name': self.context['name_latin'].split(' ')[0].upper()})
        self.context.update({'surname': self.context['name_latin'].split(' ')[1].upper()})
        # считаем 90 дней от даты заключения договора
        contract_begins = datetime.strptime(self.context['contract_begins'], '%d.%m.%Y')
        contract_ends = datetime.strftime((contract_begins + timedelta(days = 89)), '%d.%m.%Y')
        self.context.update({'contract_ends': contract_ends})
        doc = DocxTemplate("./templates/petition_tpl.docx")
        doc.render(self.context)
        filename = fr"./drivers/{self.context['name_latin']}/Ходатайство-2024 {self.context['first_name']}.docx"
        doc.save(filename)
        # если сохранение успешно - выводим об этом уведомление и увеличиваем номер ходатайства на 1 в нашем файле в облаке
        if doc.is_saved:
            print(f'Создано ходатайство на {self.context["name_cyrill"]}')
            data['petition_number'] += 1
            self.filepaths.append(fr"./drivers/{self.context['name_latin']}/Ходатайство-2024 {self.context['first_name']}.docx")
            try:
                pythoncom.CoInitialize()
                convert(fr"./drivers/{self.context['name_latin']}/Ходатайство-2024 {self.context['first_name']}.docx",
                        fr"./drivers/{self.context['name_latin']}/Ходатайство-2024 {self.context['first_name']}.pdf")
                self.filepaths.append(fr"./drivers/{self.context['name_latin']}/Ходатайство-2024 {self.context['first_name']}.pdf")
                return True
            except Exception as err:
                print('Хуйня случилась при конвертации', exc_info())


    def make_contract(self, data):
        '''делаем трудовой договор'''
        # добавляем в контекст раздробленную дату и месяц текстом
        self.context.update({'from_day': self.context['contract_begins'].split('.')[0]})
        self.context.update({'from_month': self.context['contract_begins'].split('.')[1]})
        self.context.update({'from_month_text' : self.months[self.context['from_month']]})
        self.context.update({'from_year': self.context['contract_begins'].split('.')[2]})
        if self.context['choice'] == 'recruitment_minsk':
            self.context.update({'minsk': ', трудовая деятельность осуществляется в г. Минске'})
        # проверяем, есть ли в нашем файле из облака договора, сделанные нужным числом
        if self.context['contract_begins'] in data.keys():
            # если есть - добавляем в контекст
            self.context.update({'contract_num_today': data[self.context['contract_begins']]})
        else:
            # если нет - считаем, что это первый
            self.context.update({'contract_num_today': 1})
        doc = DocxTemplate("./templates/contract_tpl.docx")
        # собираем имя будущего файла с договором
        fam = f"Трудовой договор {self.context['from_day']}_{self.context['from_month']}-{self.context['contract_num_today']} {self.context['name_latin']}"
        doc.render(self.context)
        # делаем директорию с ФИО водителя и сохраняем туда файл
        os.makedirs(fr"./drivers/{self.context['name_latin']}", exist_ok = True)
        doc.save(fr"./drivers/{self.context['name_latin']}/{fam}.docx")
        # если сохранение успешно - выводим об этом уведомление
        if doc.is_saved:
            print(f'Создан договор на {self.context["name_cyrill"]}')
            self.filepaths.append(fr"./drivers/{self.context['name_latin']}/{fam}.docx")
            # проверяем есть ли дата текущего договора в файле из облака
            if self.context['contract_begins'] in data.keys():
                # если есть - увеличиваем их количество на один
                data[self.context['contract_begins']] += 1
            else:
                # если нет - считаем, что сделали один + у следующего будет уже номер 2, который и считается в начале функции
                data[self.context['contract_begins']] = 2
            return True


    def make_notification(self):
        '''делаем уведомление'''
        doc = DocxTemplate("./templates/notification_tpl.docx")
        name = self.context['name_cyrill'].split(' ')[0]
        doc.render(self.context)
        doc.save(fr"./drivers/{self.context['name_latin']}/Уведомление УВД миграция {name}.docx")
        # если сохранение успешно - выводим об этом уведомление
        if doc.is_saved:
            self.filepaths.append(fr"./drivers/{self.context['name_latin']}/Уведомление УВД миграция {name}.docx")
            print(f"Создано уведомление УВД миграции на {self.context['name_latin']}")


    def make_dismissal_notification(self):
        '''делаем уведомление'''
        doc = DocxTemplate("./templates/dismissal_tpl.docx")
        self.context.update({'from_day': self.context['contract_begins'].split('.')[0]})
        self.context.update({'from_month': self.context['contract_begins'].split('.')[1]})
        self.context.update({'from_month_text' : self.months[self.context['from_month']]})
        self.context.update({'from_year': self.context['contract_begins'].split('.')[2]})
        self.context.update({'from_day_ends': self.context['contract_ends'].split('.')[0]})
        self.context.update({'from_month_ends': self.context['contract_ends'].split('.')[1]})
        self.context.update({'from_month_text_ends' : self.months[self.context['from_month_ends']]})
        self.context.update({'from_year_ends': self.context['contract_ends'].split('.')[2]})
        name = self.context['name_cyrill'].split(' ')[0]
        doc.render(self.context)
        os.makedirs(fr"./drivers/{self.context['name_latin']}", exist_ok = True)
        doc.save(fr"./drivers/{self.context['name_latin']}/Уведомление УВД миграция увольнение {name}.docx")
        # если сохранение успешно - выводим об этом уведомление
        if doc.is_saved:
            self.filepaths.append(fr"./drivers/{self.context['name_latin']}/Уведомление УВД миграция увольнение {name}.docx")
            print(f"Создано уведомление УВД миграции об увольнении {self.context['name_latin']}")
            return True

    def make_statement(self):
        '''делаем заявление'''
        doc = DocxTemplate("./templates/statement_tpl.docx")
        name = self.context['name_cyrill'].split(' ')[0]
        doc.render(self.context)
        doc.save(fr"./drivers/{self.context['name_latin']}/Заявление {name}.docx")
        # если сохранение успешно - выводим об этом уведомление
        if doc.is_saved:
            self.filepaths.append(fr"./drivers/{self.context['name_latin']}/Заявление {name}.docx")
            print(f"Создано заявление на {self.context['name_latin']}")

    def make_dcm_petition(self, data):
        '''делаем ходатайство в ОГИМ'''
        #открываем шаблон из корня и единственный лист в нём
        wb = load_workbook('test.xlsx')
        sheet = wb['стр.1']
        #это ссаные индексы колонок и строк в этом шаблоне. делали явно какие-то пидарасы
        indexes = {
            'name_latin': {14: range(2, 113, 4)},
            'name_cyrill': {16: range(2, 113, 4)},
            'birth_date': {18: [17, 24, 41]},
            'passport_number': {26: [10, 59, 64, 69, 74, 79, 84, 89]},
            'passport_given': {27: [12, 22, 46]},
            'passport_ends': {27: [76, 86, 110]},
            'registration_address': {35: [13, ]},
            'telephone': {37: [2, ]},
            'ogim': {5: [43, ]},
            'ogim_petition_number': {5: [6, ]},
        }
        self.context['ogim_petition_number'] = data['ogim_petition_number']
        #идём перебором по этим шаблонам, по пути подъедаем инфу из контекста
        for cat, ind in indexes.items():
            curr_val = self.context[cat]
            if cat in ['birth_date', 'passport_given', 'passport_ends']:
                curr_val = curr_val.split('.')
            for string, rows in ind.items():
                #итерируемся по пронумерованным номерам столбцов, попутно по этому индексу добавляем значения из контекста
                for num, i in enumerate(rows):
                    #если видим, что клетка для инфы - одна - вставляем туда значение из контекста целиком
                    if len(rows) == 1:
                        sheet.cell(row = string, column = i).value = curr_val
                    #иначе дрочим по одной букве
                    elif num < len(curr_val):
                        if isinstance(curr_val[num], str):
                            sheet.cell(row = string, column = i).value = curr_val[num].upper()
                        else:
                            sheet.cell(row = string, column = i).value = curr_val[num]
        #ну и сохраняем напоследок получившееся недоразумение
        try:
            name = self.context['name_cyrill'].split(' ')[0]
            wb.save(fr"./drivers/{self.context['name_latin']}/Ходатайство о рег {name}.xlsx")
            data['ogim_petition_number'] += 1
            self.filepaths.append(fr"./drivers/{self.context['name_latin']}/Ходатайство о рег {name}.xlsx")
            print(f"Создано ходатайство о регистрации на {self.context['name_latin']}")
        except:
            print('Я хз, что случилось, но случилось это при сохранении этого ебучего ходатайства в эксель', exc_info())

    def make_turkmenistan_invite(self):
        data = self.read_contracts_data()
        try:
            self.make_contract(data)
            self.make_petition(data)
            self.write_contracts_data(data)
            print('_______________________________________________')
            return True
        except Exception:
            print(exc_info())

    def make_recruitment_optional(self):
        data = self.read_contracts_data()
        try:
            self.make_contract(data)
            self.make_statement()
            self.make_notification()
            self.make_dcm_petition(data)
            self.write_contracts_data(data)
            print('_______________________________________________')
            return True
        except Exception:
            print(exc_info())

    def make_dismissal(self):
        try:
            self.make_dismissal_notification()
            print('_______________________________________________')
            return True
        except Exception:
            print(exc_info())


    def driver_exist(self):
        flag = []
        if self.context["name_latin"] in  os.listdir('drivers'):
            print('Такой водитель есть')
            for file in os.listdir(f'drivers/{self.context["name_latin"]}'):
                if file.endswith(f'{self.context["name_latin"]}.docx') and file.startswith('Трудовой договор'):
                    self.contract_filename = file
                    flag.append('Договор')
                    print(f'Договор на {self.context["name_latin"]} есть')
                if file.endswith(f'{self.context["name_latin"].split(" ")[0].upper()}.docx') and file.startswith('Ходатайство'):
                    self.petition_filename = file
                    print(f'Ходатайство на {self.context["name_latin"]} есть')
                    flag.append('Ходатайство')
        if 'Договор' in flag and 'Ходатайство' in flag:
            print(f'Нашли в системе {self.context["name_latin"]}')
        return 'Договор' in flag and 'Ходатайство' in flag

    def get_context_from_existing_driver(self):
        doc_contract = Document(f'drivers/{self.context["name_latin"]}/{self.contract_filename}')
        doc_petition = Document(f'drivers/{self.context["name_latin"]}/{self.petition_filename}')
        self.context['name_cyrill'] = doc_contract.tables[1].column_cells(2)[0].text.split('\n')[1]
        self.context['birth_date'] = doc_petition.tables[1].column_cells(0)[2].text.split('\n')[1]
        self.context['passport_number'] = doc_petition.tables[1].column_cells(0)[3].text.split('\n')[1]
        self.context['passport_given'] = doc_petition.tables[1].column_cells(2)[3].text.split('\n')[1].split('-')[0]
        self.context['passport_ends'] = doc_petition.tables[1].column_cells(2)[3].text.split('\n')[1].split('-')[1]
        print('Подсосали данные из существующих документов')
